import os
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document

SOURCE_ROOT = r"c:\Users\bryan\Documents\storymaker\Books"
EXPORT_ROOT = r"c:\Users\bryan\Documents\storymaker\Exports"
TARGET_BOOK = "Book 1 - The Iron Selection"

def process_element(paragraph, element, bold=False, italic=False):
    """
    Recursively process HTML elements and add runs to the paragraph.
    """
    if isinstance(element, NavigableString):
        text = str(element)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
    elif isinstance(element, Tag):
        is_bold = bold or element.name in ['strong', 'b']
        is_italic = italic or element.name in ['em', 'i']
        
        # Handle line breaks
        if element.name == 'br':
            paragraph.add_run('\n')
        
        for child in element.children:
            process_element(paragraph, child, bold=is_bold, italic=is_italic)

def convert_md_to_docx(md_path, docx_path):
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # Convert markdown to HTML
        html = markdown.markdown(md_text, extensions=['extra', 'nl2br'])
        soup = BeautifulSoup(html, 'html.parser')
        
        doc = Document()
        
        # We iterate over top-level elements
        for tag in soup.find_all(recursive=False):
            if tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag.name[1])
                # docx only supports levels 0-9, usually 1-9
                doc.add_heading(tag.get_text(), level=level)
            elif tag.name == 'p':
                p = doc.add_paragraph()
                for child in tag.children:
                    process_element(p, child)
            elif tag.name == 'ul':
                for li in tag.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Bullet')
                    for child in li.children:
                        process_element(p, child)
            elif tag.name == 'ol':
                for li in tag.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Number')
                    for child in li.children:
                        process_element(p, child)
            elif tag.name == 'hr':
                doc.add_paragraph('__________________________________________________')
            elif tag.name == 'pre':
                # Code block
                p = doc.add_paragraph(style='No Spacing')
                p.add_run(tag.get_text())
                p.style.font.name = 'Courier New'

        doc.save(docx_path)
        print(f"Exported: {docx_path}")
    except Exception as e:
        print(f"Failed to convert {md_path}: {e}")

def main():
    source_path = os.path.join(SOURCE_ROOT, TARGET_BOOK)
    export_path = os.path.join(EXPORT_ROOT, TARGET_BOOK)
    
    print(f"Starting export from '{source_path}' to '{export_path}'...")

    if not os.path.exists(source_path):
        print(f"Error: Source directory not found: {source_path}")
        return

    if not os.path.exists(export_path):
        os.makedirs(export_path)
        
    for root, dirs, files in os.walk(source_path):
        # Create corresponding subdirectory in Exports
        rel_path = os.path.relpath(root, source_path)
        if rel_path == '.':
            target_dir = export_path
        else:
            target_dir = os.path.join(export_path, rel_path)
        
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)
            
        for file in files:
            if file.endswith('.md'):
                md_file = os.path.join(root, file)
                docx_file = os.path.join(target_dir, file.replace('.md', '.docx'))
                convert_md_to_docx(md_file, docx_file)
    
    print("Export complete.")

if __name__ == "__main__":
    main()
